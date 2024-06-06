from docx import Document
import codecs
from docx_utils.flatten import opc_to_flat_opc
from test import somik,tomik
import MTK2
varname = 'variant12.docx'
doc = Document(varname)
opc_to_flat_opc('variant12.docx','variant12.xml')

m1 = doc.paragraphs[0].runs[-1].font.size
m2 = doc.paragraphs[0].runs[-1].font.color.rgb
m3 = doc.paragraphs[0].runs[0].font.highlight_color
print(m1,m2,m3)
for sus in range(4):
    mem = []
    for p in doc.paragraphs:
        for run in p.runs:
            #print (len(run.text),run.text,run.font.size,run.font.color.rgb,run.font.highlight_color)
            if sus == 0:
                if m1 == run.font.size:
                    mem.append("1"*len(run.text))
                else:
                    mem.append("0"*len(run.text)) 
            if sus == 1:      
                if m2 == run.font.color.rgb:
                    mem.append("1"*len(run.text))
                else: 
                    mem.append("0"*len(run.text))
            if sus == 2:
                if m3 == run.font.highlight_color:
                    mem.append("1"*len(run.text))
                else: 
                    mem.append("0"*len(run.text))
    
    out = ""
    for i in mem:
        out = out + i
    for sus in range(3):
        output = ""
        for i in range(len(out)//8):
            #print(out[0+i*8:8+i*8])
            bytes_hex = hex(int(out[0+i*8:8+i*8],2))
            #print(bytes_hex[2:4])
            if (bytes_hex != 0)&(len(bytes_hex)>3):
                if sus == 0:
                    output = output + codecs.decode(bytes.fromhex(bytes_hex[2:4]), encoding='koi8-r')
                if sus == 1:
                    output = output + codecs.decode(bytes.fromhex(bytes_hex[2:4]), encoding='cp866')
                if sus == 2:
                    output = output + codecs.decode(bytes.fromhex(bytes_hex[2:4]), encoding='cp1251')
        #print(out)
        print(output)
    output = MTK2.MTK2_decode(out)
    print(output)
        
out = tomik(varname.replace(".docx",".xml"))

#print(out)
for sus in range(3):
        output = ""
        for i in range(len(out)//8):
            #print(out[0+i*8:8+i*8])
            bytes_hex = hex(int(out[0+i*8:8+i*8],2))
            #print(bytes_hex[2:4])
            if (bytes_hex != 0)&(len(bytes_hex)>3):
                if sus == 0:
                    output = output + codecs.decode(bytes.fromhex(bytes_hex[2:4]), encoding='koi8-r')
                if sus == 1:
                    output = output + codecs.decode(bytes.fromhex(bytes_hex[2:4]), encoding='cp866')
                if sus == 2:
                    output = output + codecs.decode(bytes.fromhex(bytes_hex[2:4]), encoding='cp1251')
        print(output)
output = MTK2.MTK2_decode(out)
print(output)
#print(huh.decode('koi8-r'))
#print(huh.decode('cp866'))
#print(huh.decode('cp1251'))
out = somik(varname.replace(".docx",".xml"))
for sus in range(3):
        output = ""
        for i in range(len(out)//8):
            #print(out[0+i*8:8+i*8])
            bytes_hex = hex(int(out[0+i*8:8+i*8],2))
            #print(bytes_hex[2:4])
            if (bytes_hex != 0)&(len(bytes_hex)>3):
                if sus == 0:
                    output = output + codecs.decode(bytes.fromhex(bytes_hex[2:4]), encoding='koi8-r')
                if sus == 1:
                    output = output + codecs.decode(bytes.fromhex(bytes_hex[2:4]), encoding='cp866')
                if sus == 2:
                    output = output + codecs.decode(bytes.fromhex(bytes_hex[2:4]), encoding='cp1251')
        print(output)
output = MTK2.MTK2_decode(out)
print(output)
doc.save('test.docx')


#lol = "сос"
#print(type(lol))
#koi = lol.encode(encoding='koi8-r',errors='ignorea')
#print(koi)
#print(koi.decode('koi8-r'))
